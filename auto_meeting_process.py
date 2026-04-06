import os
import time
import shutil
import smtplib
import logging
import re
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

import torch
from transformers import pipeline
from docx import Document

# GPU 互斥鎖（與 auto_youtube_whisper.py 共用）
from gpu_lock import acquire_gpu_lock, release_gpu_lock

# faster-whisper（CTranslate2 加速）
try:
    from faster_whisper import WhisperModel as FasterWhisperModel
    FASTER_WHISPER_AVAILABLE = True
except ImportError:
    FASTER_WHISPER_AVAILABLE = False

# ================= 設定區域 =================
# 監控與輸出路徑
WATCH_DIR = "/mnt/nas/Whisper_auto_rum/月會聽打"
PROCESSED_DIR = os.path.join(WATCH_DIR, "processed")
OUTPUT_DIR = os.path.join(WATCH_DIR, "output")

# Email 設定
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
EMAIL_USER = "bbudaedu@gmail.com"
EMAIL_PASS = "khwvartcqbgztwgr"
EMAIL_TO = ["jackyfang@budaedu.org", "tyguo@budaedu.org", "zrwang@budaedu.org"]

# AI 模型設定
PUNCTUATION_MODEL = "oliverguhr/fullstop-punctuation-multilang-large"
WHISPER_MODEL = "large-v2"
WHISPER_LANG = "zh"
WHISPER_PROMPT = "會議記錄 "

# ==========================================

ENABLE_DIARIZATION = os.getenv("ENABLE_DIARIZATION", "false").lower() in {"1", "true", "yes"}

# 設定日誌
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("meeting_process.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# 行程層級模型快取（同一次執行多檔時只載入一次）
_whisper_model_cache: dict = {}


def setup_directories():
    for d in [PROCESSED_DIR, OUTPUT_DIR]:
        os.makedirs(d, exist_ok=True)


def send_email(subject, body, attachment_path):
    msg = MIMEMultipart()
    msg['From'] = EMAIL_USER
    msg['To'] = ", ".join(EMAIL_TO)
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    if attachment_path and os.path.exists(attachment_path):
        filename = os.path.basename(attachment_path)
        with open(attachment_path, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())

        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            "attachment",
            filename=('utf-8', '', filename)
        )
        msg.attach(part)

    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(EMAIL_USER, EMAIL_PASS)
        server.sendmail(EMAIL_USER, EMAIL_TO, msg.as_string())
        server.quit()
        logger.info(f"Email 已發送至 {EMAIL_TO}")
    except Exception as e:
        logger.error(f"Email 發送失敗: {e}")


def _get_whisper_model(model_name: str) -> "FasterWhisperModel":
    """取得（或建立）快取的 faster-whisper WhisperModel 實例"""
    if model_name not in _whisper_model_cache:
        logger.info(f"載入 faster-whisper 模型: {model_name} (float16, CUDA)")
        _whisper_model_cache[model_name] = FasterWhisperModel(
            model_name,
            device="cuda",
            compute_type="float16",
        )
    return _whisper_model_cache[model_name]


def run_whisper(filepath):
    """使用 faster-whisper Python API 進行轉錄"""
    logger.info(f"正在執行 faster-whisper 轉錄: {os.path.basename(filepath)}")

    if not FASTER_WHISPER_AVAILABLE:
        logger.error("faster-whisper 未安裝，請執行: pip install faster-whisper")
        return None

    try:
        model = _get_whisper_model(WHISPER_MODEL)
        segments_gen, info = model.transcribe(
            filepath,
            language=WHISPER_LANG,
            initial_prompt=WHISPER_PROMPT,
            vad_filter=True,
            beam_size=5,
        )
        segments = list(segments_gen)
        logger.info(
            f"辨識完成，共 {len(segments)} 段，"
            f"語言: {info.language} (機率: {info.language_probability:.2f})"
        )

        # 組合純文字
        raw_text = "".join(seg.text.strip() for seg in segments)
        return {"raw_text": raw_text, "segments": segments}

    except Exception as e:
        logger.error(f"faster-whisper 辨識失敗: {e}")
        return None


def apply_punctuation(text, pipe):
    """使用模型修復標點符號"""
    logger.info("正在進行標點符號修復...")

    # 清理並切分文本
    text = text.replace("\n", "").replace("\r", "").replace(" ", "")
    max_chunk_chars = 250
    chunks = [text[i:i + max_chunk_chars] for i in range(0, len(text), max_chunk_chars)]

    full_punctuated_text = ""
    for chunk in chunks:
        result = pipe(chunk)
        for item in result:
            full_punctuated_text += item['word']
            tag = item['entity_group']
            if tag != 'O' and tag != '0':
                full_punctuated_text += tag

    # 格式化清理
    final = full_punctuated_text.replace("##", "")
    final = re.sub(r'([\u4e00-\u9fa5])\s+([\u4e00-\u9fa5])', r'\1\2', final)
    final = final.replace(",", "，").replace(".", "。").replace("?", "？").replace("!", "！").replace(":", "：")
    for punc in "，。？！：":
        final = final.replace(" " + punc, punc)

    return final


def create_docx(text, output_path):
    """將文字轉換為 DOCX"""
    doc = Document()
    doc.add_heading('會議記錄聽打結果', 0)

    # 簡單分段
    paragraphs = text.split("。")
    for p in paragraphs:
        if p.strip():
            doc.add_paragraph(p.strip() + "。")

    doc.save(output_path)
    logger.info(f"已建立 Word 檔: {output_path}")


def main():
    setup_directories()
    logger.info("系統啟動完成，進入按需載入 (Lazy Loading) 監控模式...")

    while True:
        try:
            if not os.path.exists(WATCH_DIR):
                time.sleep(10)
                continue

            # 監控音檔 (mp3, wav, m4a)
            files = [f for f in os.listdir(WATCH_DIR) if f.lower().endswith(('.mp3', '.wav', '.m4a'))]
            
            # 過濾掉尚未寫入完成的檔案
            pending_files = []
            for file in files:
                filepath = os.path.join(WATCH_DIR, file)
                if os.path.getsize(filepath) > 0:
                    pending_files.append(file)

            if not pending_files:
                time.sleep(10)
                continue

            logger.info(f"發現 {len(pending_files)} 個音檔等待處理")

            # === GPU 互斥鎖：確保不會與 auto_youtube_whisper 同時使用 GPU ===
            gpu_fd = acquire_gpu_lock()
            if gpu_fd is None:
                logger.info("GPU 被其他行程佔用中，本輪跳過，10 秒後重試...")
                time.sleep(10)
                continue

            punctuator = None
            try:
                import gc
                # 只有在確定拿到鎖且有檔案時，才載入模型
                logger.info("正在載入標點與 Whisper 模型至 GPU...")
                device = 0 if torch.cuda.is_available() else -1
                punctuator = pipeline("token-classification", model=PUNCTUATION_MODEL, device=device, aggregation_strategy="simple")

                for file in pending_files:
                    filepath = os.path.join(WATCH_DIR, file)
                    logger.info(f"開始處理: {file}")

                    # 1. 執行 faster-whisper 轉錄
                    whisper_result = run_whisper(filepath)
                    if not whisper_result:
                        continue

                    raw_text = whisper_result["raw_text"]
                    segments = whisper_result["segments"]

                    if ENABLE_DIARIZATION:
                        try:
                            from pipeline.diarization import run_diarization

                            diarization = run_diarization(filepath)
                            speaker_paragraphs = []
                            current_label = None
                            current_text = []

                            for seg in segments:
                                midpoint = (seg.start + seg.end) / 2.0
                                label = None
                                for start, end, speaker in diarization:
                                    if start <= midpoint <= end:
                                        label = speaker
                                        break

                                if label != current_label and current_text:
                                    speaker_paragraphs.append((current_label, "".join(current_text)))
                                    current_text = []

                                current_label = label
                                current_text.append(seg.text.strip())

                            if current_text:
                                speaker_paragraphs.append((current_label, "".join(current_text)))

                            labeled_blocks = []
                            for label, block in speaker_paragraphs:
                                punctuated = apply_punctuation(block, punctuator)
                                if label:
                                    labeled_blocks.append(f"{label}: {punctuated}")
                                else:
                                    labeled_blocks.append(punctuated)

                            final_text = "\n".join(labeled_blocks)
                        except Exception as e:
                            logger.error(f"Diarization 失敗，改用無標籤輸出: {e}")
                            final_text = apply_punctuation(raw_text, punctuator)
                    else:
                        # 2. 執行標點修復
                        final_text = apply_punctuation(raw_text, punctuator)

                    # 3. 產生 DOCX
                    base_name = os.path.splitext(file)[0]
                    docx_filename = f"{base_name}_punctuated.docx"
                    docx_path = os.path.join(OUTPUT_DIR, docx_filename)
                    create_docx(final_text, docx_path)

                    # 4. 發送 Email
                    subject = f"【會議聽打完成】{file}"
                    body = f"檔案 {file} 已完成轉錄與標點修復。\n\n請查收附件 Word 檔。\n\n系統自動發送"
                    send_email(subject, body, docx_path)

                    # 5. 清理與歸檔
                    shutil.move(filepath, os.path.join(PROCESSED_DIR, file))
                    logger.info(f"任務完成: {file}")

            finally:
                # 處理完這批檔案後，釋放模型與記憶體
                logger.info("釋放 AI 模型與 GPU 記憶體，回歸待命狀態...")
                if punctuator:
                    del punctuator
                
                # 清空 faster-whisper 的模型快取
                global _whisper_model_cache
                _whisper_model_cache.clear()
                
                # 強制回收資源
                gc.collect()
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()

                # 釋放 GPU 鎖
                release_gpu_lock(gpu_fd)

            time.sleep(10)

        except Exception as e:
            logger.error(f"主迴圈錯誤: {e}")
            time.sleep(10)


if __name__ == "__main__":
    main()
