import sys
import os
import math
import re
import pandas as pd
import openpyxl
import difflib
import docx
from collections import Counter
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLineEdit, QTextEdit, QFileDialog, QLabel, QMessageBox,
    QTabWidget
)
from PyQt6.QtCore import QThread, QObject, pyqtSignal, Qt
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter

# ===================================================================
#
#  第一部分：核心處理邏輯
#
# ===================================================================

def simple_autofit_columns(output_path, log_callback):
    try:
        wb = openpyxl.load_workbook(output_path)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for col in ws.columns:
                max_length = 0
                column_letter = get_column_letter(col[0].column)
                header_value = ws[1][col[0].column - 1].value
                if header_value:
                    max_length = len(str(header_value)) * 1.2

                for cell in col:
                    try:
                        if cell.value:
                            cell_len = sum(2 if '\u4e00' <= char <= '\u9fff' else 1 for char in str(cell.value))
                            if cell_len > max_length:
                                max_length = cell_len
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column_letter].width = adjusted_width
        wb.save(output_path)
        log_callback("✅ 已自動調整欄位寬度。")
    except Exception as e:
        log_callback(f"⚠️ 自動調整欄寬時發生錯誤: {e}")

def run_dictionary_correction(df, dict_filepath, source_col_name, output_col_name, log_callback):
    sheet_name = '校對字典'
    error_col = '錯誤詞'
    correct_col = '正確詞'
    try:
        dict_df = pd.read_excel(dict_filepath, sheet_name=sheet_name)
    except Exception:
        log_callback(f"[警告] 找不到名為 '{sheet_name}' 的工作表，將嘗試讀取第一個工作表。")
        dict_df = pd.read_excel(dict_filepath, sheet_name=0)
    
    dict_df.columns = [str(c).strip() for c in dict_df.columns]
    if error_col not in dict_df.columns or correct_col not in dict_df.columns:
        raise ValueError(f"字典檔案缺少 '{error_col}' 或 '{correct_col}' 欄位。")
    
    correction_dict = dict(zip(dict_df[error_col].dropna().astype(str), dict_df[correct_col].dropna().astype(str)))
    log_callback(f"  - 成功從 '{os.path.basename(dict_filepath)}' 載入 {len(correction_dict)} 條校對規則。")
    
    df[source_col_name] = df[source_col_name].fillna('').astype(str)
    corrected_texts, usage_stats, modified_row_numbers = [], Counter(), []
    
    log_callback("  - 正在逐行應用校對規則...")
    for index, row in df.iterrows():
        original_text = row[source_col_name]
        corrected_text = original_text
        is_changed = False
        for error, correct in correction_dict.items():
            if error in corrected_text:
                count = corrected_text.count(error)
                corrected_text = corrected_text.replace(error, correct)
                usage_stats[f"'{error}' -> '{correct}'"] += count
                is_changed = True
        if is_changed:
            modified_row_numbers.append(index + 2)
        corrected_texts.append(corrected_text)
    
    df[output_col_name] = corrected_texts
    log_callback("  - 校對規則應用完畢。")

    stats = {
        "total_rows": len(df), "changed_rows": len(modified_row_numbers),
        "total_replacements": sum(usage_stats.values()), "usage_stats": usage_stats,
        "modified_row_numbers": modified_row_numbers
    }
    return df, stats

def run_aligner_task(file1_path, file2_path, dict_filepath, output_path, log_callback):
    def read_file_lines(filename):
        log_callback(f"正在讀取檔案: {os.path.basename(filename)}...")
        lines = [line.strip() for line in open(filename, 'r', encoding='utf-8') if line.strip()]
        log_callback(f"✅ 成功讀取 {len(lines)} 行。")
        return lines
    try:
        lines1, lines2 = read_file_lines(file1_path), read_file_lines(file2_path)
        initial_lines1_count = len(lines1)
        
        log_callback("\n正在進行序列對齊...")
        aligned1, aligned2 = align_sequences(lines1, lines2)
        log_callback("✅ 對齊完成！")
        
        log_callback("\n⚙️ 開始進行Excel後處理...")
        # 統計資訊
        missing_count = 0
        corrected_count = 0
        missing_list = [] # (序號, 內容)
        
        # 1. 識別遺漏句 (Gemini 為空之處)
        for idx, (w, g) in enumerate(zip(aligned1, aligned2)):
            if w != '' and g == '':
                missing_count += 1
                missing_list.append((idx + 1, w))
        
        # 2. 識別校對句 (Whisper 與 Gemini 不同且 Gemini 不為空之處)
        gemini_omissions = []
        gemini_corrections = []
        for w, g in zip(aligned1, aligned2):
            if w != '' and g != '' and w != g:
                corrected_count += 1
                gemini_omissions.append("")
                gemini_corrections.append(g)
            elif w != '' and g == '':
                gemini_omissions.append(w)
                gemini_corrections.append("")
            else:
                gemini_omissions.append("")
                gemini_corrections.append("")

        df = pd.DataFrame({
            'whisper': aligned1, 
            'gemini': aligned2,
            'Gemini遺漏句': gemini_omissions,
            'Gemini校對句': gemini_corrections
        })
        df = df[df['whisper'] != ''].copy()
        df.loc[df['gemini'] == '', 'gemini'] = df['whisper']
        log_callback(f"  - 規則處理完成。")
        log_callback(f"  - 發現 Gemini 遺漏句: {missing_count} 句")
        log_callback(f"  - 發現 Gemini 校對句: {corrected_count} 句")

        final_lines_count = len(df)
        log_callback("\n📊 進行句數檢查...")
        log_callback(f"  - 原始句數: {initial_lines1_count}, 處理後句數: {final_lines_count}")
        if initial_lines1_count != final_lines_count:
            log_callback(f"  - ⚠️ 警告: 句數不符！")
        else:
            log_callback("  - ✅ 結果: 句數相符。")

        if dict_filepath and os.path.exists(dict_filepath):
            log_callback("\n=== 開始自動校對程序 ===")
            df, stats = run_dictionary_correction(df, dict_filepath, 'gemini', '校對字典', log_callback)
            
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='校對文本', index=False)
                # 原始統計
                summary_df = pd.DataFrame([("總處理行數", stats['total_rows']), ("修改行數", stats['changed_rows']), ("總替換次數", stats['total_replacements'])])
                usage_df = pd.DataFrame(stats['usage_stats'].most_common(), columns=["替換規則", "次數"])
                modified_rows_df = pd.DataFrame(stats['modified_row_numbers'], columns=["被修改行號(Excel)"])
                summary_df.to_excel(writer, sheet_name='校對字典統計', index=False, header=False, startrow=0)
                usage_df.to_excel(writer, sheet_name='校對字典統計', index=False, startrow=len(summary_df) + 2)
                modified_rows_df.to_excel(writer, sheet_name='校對字典統計', index=False, startrow=len(summary_df) + len(usage_df) + 4)
                
                # Gemini 專屬統計
                gemini_stats_df = pd.DataFrame([
                    ("Gemini 總句數", initial_lines1_count),
                    ("Gemini 遺漏句數", missing_count),
                    ("Gemini 校對句數", corrected_count),
                    ("遺漏率", f"{(missing_count/initial_lines1_count*100):.2f}%" if initial_lines1_count > 0 else "0%"),
                    ("校對率", f"{(corrected_count/initial_lines1_count*100):.2f}%" if initial_lines1_count > 0 else "0%")
                ])
                missing_detail_df = pd.DataFrame(missing_list, columns=["原始序號", "遺漏內容 (Whisper)"])
                gemini_stats_df.to_excel(writer, sheet_name='Gemini校對統計', index=False, header=False, startrow=0)
                missing_detail_df.to_excel(writer, sheet_name='Gemini校對統計', index=False, startrow=len(gemini_stats_df) + 2)
                
            final_message = "對齊與校對成功！"
        else:
            log_callback("\n未提供字典，僅執行對齊。")
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='校對文本')
                # Gemini 專屬統計 (即使沒有字典也要輸出)
                gemini_stats_df = pd.DataFrame([
                    ("Gemini 總句數", initial_lines1_count),
                    ("Gemini 遺漏句數", missing_count),
                    ("Gemini 校對句數", corrected_count),
                    ("遺漏率", f"{(missing_count/initial_lines1_count*100):.2f}%" if initial_lines1_count > 0 else "0%"),
                    ("校對率", f"{(corrected_count/initial_lines1_count*100):.2f}%" if initial_lines1_count > 0 else "0%")
                ])
                missing_detail_df = pd.DataFrame(missing_list, columns=["原始序號", "遺漏內容 (Whisper)"])
                gemini_stats_df.to_excel(writer, sheet_name='Gemini校對統計', index=False, header=False, startrow=0)
                missing_detail_df.to_excel(writer, sheet_name='Gemini校對統計', index=False, startrow=len(gemini_stats_df) + 2)
            final_message = "對齊成功！"
        
        log_callback("\n🎨 正在自動調整欄位寬度...")
        simple_autofit_columns(output_path, log_callback)
        log_callback("✅ 所有處理完成！")
        
        return f"{final_message}\n結果已儲存至: {os.path.basename(output_path)}"
    except Exception as e: 
        raise e

def run_processor_task(base_dir, file_base_name, log_callback):
    def srt_time_to_seconds(time_str):
        if not time_str or time_str.count(':') != 2 or ',' not in time_str: return 0.0
        parts = time_str.split(','); h_m_s = parts[0].split(':')
        return int(h_m_s[0])*3600 + int(h_m_s[1])*60 + int(h_m_s[2]) + int(parts[1])/1000.0
    def seconds_to_display_time(total_seconds):
        if total_seconds < 0: total_seconds = 0
        hours, rem = divmod(total_seconds, 3600); minutes, seconds_float = divmod(rem, 60); seconds = int(seconds_float)
        milliseconds = int((seconds_float - seconds) * 1000)
        return f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d},{milliseconds:03d}"
    try:
        excel_path = os.path.join(base_dir, f"{file_base_name}.xlsx"); srt_path = os.path.join(base_dir, f"{file_base_name}.srt")
        log_callback(f"腳本將處理以下檔案:\n  - Excel: {excel_path}\n  - SRT: {srt_path}")
        sheet_names_to_try = ["文本校對", "校對文本"]; df = None; loaded_sheet_name = ""
        try:
            workbook = openpyxl.load_workbook(excel_path, read_only=True); available_sheets = workbook.sheetnames; workbook.close()
            for name in sheet_names_to_try:
                if name in available_sheets:
                    log_callback(f"找到工作表: '{name}'。"); df = pd.read_excel(excel_path, header=0, sheet_name=name); loaded_sheet_name = name; break
            if df is None: raise ValueError(f"找不到 '{sheet_names_to_try[0]}' 或 '{sheet_names_to_try[1]}' 工作表。")
        except Exception as e: raise ValueError(f"讀取 Excel 失敗: {e}")
        
        if len(df.columns) < 3:
            log_callback(f"警告: 在 '{loaded_sheet_name}' 中找不到C欄，將從B欄讀取資料。")
            corrected_texts = df.iloc[:, 1].astype(str).str.strip().tolist()
        else:
            log_callback(f"正在從 '{loaded_sheet_name}' 的 C 欄讀取資料。")
            corrected_texts = df.iloc[:, 2].astype(str).str.strip().tolist()
        log_callback(f"✅ 成功讀取 {len(corrected_texts)} 筆文本。")

        time_data = [];
        with open(srt_path, 'r', encoding='utf-8') as f:
            for line in f:
                if '-->' in line: start_str, end_str = [t.strip() for t in line.split('-->')]; time_data.append({'start_str': start_str, 'end_str': end_str})
        log_callback(f"✅ 成功從 SRT 檔案讀取 {len(time_data)} 筆時間軸資料。")
        if len(corrected_texts) != len(time_data): raise ValueError(f"資料筆數不匹配。文本: {len(corrected_texts)}, 時間軸: {len(time_data)}。")
        if not corrected_texts: raise ValueError("沒有提取到任何資料。")
        
        items_with_time = []
        for i in range(len(time_data)):
            items_with_time.append({
                'num': i + 1,
                'text': f"{i + 1} {corrected_texts[i]}",
                'start_s': srt_time_to_seconds(time_data[i]['start_str']),
                'end_s': srt_time_to_seconds(time_data[i]['end_str']),
                'start_str': time_data[i]['start_str'],
                'end_str': time_data[i]['end_str']
            })

        total_audio_duration_seconds = items_with_time[-1]['end_s'] if items_with_time else 0
        segment_duration_seconds = 30 * 60
        total_parts = math.ceil(total_audio_duration_seconds / segment_duration_seconds) if total_audio_duration_seconds > 0 else 1
        log_callback(f"總音訊時長: {seconds_to_display_time(total_audio_duration_seconds)}。預計分割成 {total_parts} 個部分。")
        
        rows_for_excel_output, part_number, current_part_items = [], 1, []
        for item in items_with_time:
            segment_boundary_s = part_number * segment_duration_seconds
            if item['start_s'] >= segment_boundary_s and current_part_items:
                log_callback(f"\n--- 正在生成第 {part_number} 部分 ---")
                start_sentence_num, end_sentence_num = current_part_items[0]['num'], current_part_items[-1]['num']
                header_line1 = f"{file_base_name} Part {part_number} of {total_parts}"
                header_line2 = f"{start_sentence_num}-{end_sentence_num}句"
                header_line3 = f"{current_part_items[0]['start_str']} --> {current_part_items[0]['end_str']}"
                log_callback(f"  [Info] 第 {part_number} 部分時間戳: {header_line3}")
                rows_for_excel_output.extend([[header_line1], [header_line2], [header_line3]])
                for part_item in current_part_items: rows_for_excel_output.append([part_item['text']])
                rows_for_excel_output.append([""]); part_number += 1; current_part_items = []
            current_part_items.append(item)
            
        if current_part_items:
            log_callback(f"\n--- 正在生成最後一部分 (第 {part_number} 部分) ---")
            start_sentence_num, end_sentence_num = current_part_items[0]['num'], current_part_items[-1]['num']
            header_line1 = f"{file_base_name} Part {part_number} of {total_parts}"
            header_line2 = f"{start_sentence_num}-{end_sentence_num}句"
            header_line3 = f"{current_part_items[0]['start_str']} --> {current_part_items[0]['end_str']}"
            log_callback(f"  [Info] 第 {part_number} 部分時間戳: {header_line3}")
            rows_for_excel_output.extend([[header_line1], [header_line2], [header_line3]])
            for part_item in current_part_items: rows_for_excel_output.append([part_item['text']])

        if rows_for_excel_output:
            log_callback(f"\n✍️ 正在將 '給學員校對' 工作表寫入 Excel 檔案中...")
            workbook = openpyxl.load_workbook(excel_path); sheet_name = "給學員校對"
            if sheet_name in workbook.sheetnames: del workbook[sheet_name]
            worksheet = workbook.create_sheet(title=sheet_name)
            for row_data in rows_for_excel_output: worksheet.append(row_data)
            workbook.save(excel_path)
            log_callback(f"✅ 成功寫入 '{sheet_name}' 工作表。")

        return "處理成功！"
    except Exception as e:
        raise e

def run_converter_task(srt_path, log_callback):
    excel_file_path = os.path.splitext(srt_path)[0] + '.xlsx'
    if not os.path.exists(excel_file_path):
        raise FileNotFoundError(f"找不到對應的 Excel: {os.path.basename(excel_file_path)}")
    
    with open(srt_path, 'r', encoding='utf-8-sig') as f: content = f.read()
    data = re.findall(r'(\d+)\n(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})\n(.*?)(?=\n\n|\Z)', content, re.DOTALL)
    df = pd.DataFrame(data, columns=['序號', '開始時間', '結束時間', '文字'])
    df['文字'] = df['文字'].str.replace('\n', ' ', regex=False)

    with pd.ExcelWriter(excel_file_path, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
        df.to_excel(writer, sheet_name='時間軸', index=False)
    
    log_callback(f"✅ 成功將'時間軸'工作表加入到 {os.path.basename(excel_file_path)}")
    return "轉換成功！"

def get_similarity_ratio(s1, s2):
    return difflib.SequenceMatcher(None, s1, s2, autojunk=False).ratio()

def align_sequences(seq1, seq2, gap_penalty=-0.5):
    n, m = len(seq1), len(seq2)
    dp = [[(0, '')] * (m + 1) for _ in range(n + 1)]
    for i in range(1, n + 1): dp[i][0] = (i * gap_penalty, 'up')
    for j in range(1, m + 1): dp[0][j] = (j * gap_penalty, 'left')
    
    for i in range(1, n + 1):
        for j in range(1, m + 1):
            match = dp[i-1][j-1][0] + (2 * get_similarity_ratio(seq1[i-1], seq2[j-1]) - 1)
            scores = {'diag': match, 'up': dp[i-1][j][0] + gap_penalty, 'left': dp[i][j-1][0] + gap_penalty}
            best_move = max(scores, key=scores.get)
            dp[i][j] = (scores[best_move], best_move)

    align1, align2 = [], []
    i, j = n, m
    while i > 0 or j > 0:
        move = dp[i][j][1]
        if move == 'diag': align1.append(seq1[i-1]); align2.append(seq2[j-1]); i-=1; j-=1
        elif move == 'up': align1.append(seq1[i-1]); align2.append(""); i-=1
        else: align1.append(""); align2.append(seq2[j-1]); j-=1
    return align1[::-1], align2[::-1]

# --- 【功能4. 核心函式修改處】 ---
def read_text_from_excel_for_docx(file_path, base_name, log_callback):
    try:
        workbook = openpyxl.load_workbook(file_path, read_only=True)
        target_sheet = None
        # Assuming the sheet is named "給學員校對" as per function 3
        sheet_name_to_find = "給學員校對"
        if sheet_name_to_find in workbook.sheetnames:
            target_sheet = workbook[sheet_name_to_find]
            log_callback(f"✓ 找到工作表: '{sheet_name_to_find}'")
        else:
            raise ValueError(f"錯誤：在檔案中找不到 '{sheet_name_to_find}' 工作表。")
        
        text_lines = []
        sentence_count = 0
        for row in target_sheet.iter_rows(min_row=1, max_col=1, values_only=True):
            cell_value = row[0]
            if cell_value is not None:
                line_str = str(cell_value)
                text_lines.append(line_str)
                if re.match(r"^\d+ ", line_str.strip()):
                    sentence_count += 1
        
        log_callback(f"✓ 在 Excel 中找到 {sentence_count} 個句子。")
        return "\n".join(text_lines), sentence_count
    except Exception as e:
        raise Exception(f"讀取 Excel 檔案時發生錯誤: {e}")

def create_formatted_docx(text_content, output_filename, base_name, expected_sentence_count):
    # 【核心修復】在此函式開頭建立文件物件
    document = docx.Document()
    
    style = document.styles['Normal']; font = style.font; p_format = style.paragraph_format
    font.name = '標楷體'; style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體'); font.size = Pt(16)
    p_format.space_before = Pt(0); p_format.space_after = Pt(0)
    p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY; p_format.line_spacing = Pt(24)

    section = document.sections[0]
    section.top_margin = Cm(1.27); section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27); section.right_margin = Cm(1.27)
    section.header_distance = Cm(0.6); section.footer_distance = Cm(0.6)
    
    footer = section.footer
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run(); fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin'); instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'; fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end'); run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    
    sectPr = section._sectPr
    cols = OxmlElement('w:cols'); cols.set(qn('w:num'), '2'); cols.set(qn('w:space'), str(Cm(1).twips)); cols.set(qn('w:sep'), 'true'); sectPr.append(cols)
    
    splitter = f'{base_name} Part '
    parts_raw = text_content.split(splitter)[1:]
    
    total_written_sentences = 0
    is_first_part = True
    for part_content in parts_raw:
        if not is_first_part:
            document.add_section(WD_SECTION_START.NEW_PAGE)
            new_section = document.sections[-1]
            new_section.top_margin = Cm(1.27); new_section.bottom_margin = Cm(1.27); new_section.left_margin = Cm(1.27); new_section.right_margin = Cm(1.27)
            new_section.header_distance = Cm(0.6); new_section.footer_distance = Cm(0.6)
            sectPr = new_section._sectPr; cols = OxmlElement('w:cols'); cols.set(qn('w:num'), '2'); cols.set(qn('w:space'), str(Cm(1).twips)); cols.set(qn('w:sep'), 'true'); sectPr.append(cols)

        lines = part_content.strip().split('\n')
        header_line = splitter + lines[0]
        doc_para = document.add_paragraph(); doc_para.text = header_line

        for line in lines[1:]:
            doc_para = document.add_paragraph(); doc_para.text = line
            if re.match(r"^\d+ ", line.strip()):
                total_written_sentences += 1
        is_first_part = False
        
    if total_written_sentences != expected_sentence_count:
        error_msg = (f"‼️ 句子數量不匹配！\n\n"
                     f"從 Excel 讀取了 {expected_sentence_count} 句，\n"
                     f"但準備寫入 Word 的只有 {total_written_sentences} 句。\n\n"
                     f"請檢查原始文本內容。")
        raise Exception(error_msg)
        
    document.save(output_filename)

def run_docx_formatter_task(file_path, log_callback):
    try:
        log_callback("\n開始處理文件...")
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        text_content, expected_count = read_text_from_excel_for_docx(file_path, base_name, log_callback)
        log_callback(f"✅ Excel 內容讀取成功，預計處理 {expected_count} 個句子。")
        output_filename = os.path.join(os.path.dirname(file_path), f"{base_name}給學員校對.docx")
        create_formatted_docx(text_content, output_filename, base_name, expected_count)
        log_callback("✓ 句子數量檢查通過。")
        log_callback(f"✓ 文件已成功建立！儲存路徑: {output_filename}")
        return f"文件已成功建立！\n兩欄式版面配置已自動套用。\n\n儲存於:\n{output_filename}"
    except Exception as e:
        log_callback(f"❌ 處理過程中發生錯誤: {e}"); raise e

# ===================================================================
#
#  第二部分：背景執行緒 (Worker)
#
# ===================================================================

class Worker(QObject):
    finished = pyqtSignal(str); progress = pyqtSignal(str)
    def __init__(self, task, *args):
        super().__init__(); self.task = task; self.args = args
    def run(self):
        try: self.finished.emit(self.task(*self.args, self.progress.emit))
        except Exception as e: self.finished.emit(f"處理失敗: {e}")

# ===================================================================
#
#  第三部分：UI 介面定義
#
# ===================================================================

class BaseUI(QWidget):
    def __init__(self):
        super().__init__(); self.thread, self.worker = None, None
    def start_task(self, task, *args):
        self.set_buttons_enabled(False); self.log_display.clear()
        self.thread = QThread(); self.worker = Worker(task, *args)
        self.worker.moveToThread(self.thread)
        self.thread.started.connect(self.worker.run)
        self.worker.finished.connect(self.on_finished)
        self.worker.progress.connect(self.log_display.append)
        self.thread.start()
    def on_finished(self, message):
        self.log_display.append(f"\n🎉 {message}")
        self.set_buttons_enabled(True)
        if "失敗" in message: QMessageBox.critical(self, "失敗", message)
        else: QMessageBox.information(self, "完成", message)
        if self.thread: self.thread.quit(); self.thread.wait()
    def set_buttons_enabled(self, enabled):
        for btn in self.findChildren(QPushButton): btn.setEnabled(enabled)

class AlignerUI(BaseUI):
    def __init__(self): super().__init__(); self.init_ui()
    def init_ui(self):
        layout = QVBoxLayout(self)
        self.file1_path_edit = self._create_file_selector(layout, "選擇檔案 1 (例如 Whisper.txt)", "Text Files (*.txt)")
        self.file2_path_edit = self._create_file_selector(layout, "選擇檔案 2 (例如 Gemini.txt)", "Text Files (*.txt)")
        self.dict_path_edit = self._create_file_selector(layout, "選擇檔案 3 (校對字典 Excel，可選)", "Excel Files (*.xlsx *.xls)")
        self.start_button = QPushButton("開始對齊與校對"); self.start_button.setEnabled(False)
        self.start_button.clicked.connect(self.start_alignment)
        layout.addWidget(self.start_button); layout.addWidget(QLabel("處理日誌："))
        self.log_display = QTextEdit(); self.log_display.setReadOnly(True); layout.addWidget(self.log_display)
    def _create_file_selector(self, layout, label, file_filter):
        layout.addWidget(QLabel(label)); h_layout = QHBoxLayout(); line_edit = QLineEdit(); line_edit.setReadOnly(True)
        browse_button = QPushButton("瀏覽..."); browse_button.clicked.connect(lambda: self._select_file(line_edit, file_filter))
        h_layout.addWidget(line_edit); h_layout.addWidget(browse_button); layout.addLayout(h_layout); return line_edit
    def _select_file(self, line_edit, file_filter):
        file_name, _ = QFileDialog.getOpenFileName(self, "選擇檔案", "", file_filter)
        if file_name:
            line_edit.setText(file_name); self.log_display.append(f"已選擇: {os.path.basename(file_name)}")
            if self.file1_path_edit.text() and self.file2_path_edit.text(): self.start_button.setEnabled(True)
    def start_alignment(self):
        file1, file2, dict_file = self.file1_path_edit.text(), self.file2_path_edit.text(), self.dict_path_edit.text()
        output_path = os.path.splitext(file1)[0] + '.xlsx'
        msg = f"結果將儲存為：\n{output_path}\n" + ("並進行字典校對。\n" if dict_file else "") + "確定嗎？"
        if QMessageBox.question(self, '確認執行', msg) == QMessageBox.StandardButton.Yes:
            self.start_task(run_aligner_task, file1, file2, dict_file, output_path)

class ProcessorUI(BaseUI):
    def __init__(self): super().__init__(); self.init_ui()
    def init_ui(self):
        layout = QVBoxLayout(self); h_layout = QHBoxLayout()
        self.folder_path_edit = QLineEdit(); self.folder_path_edit.setReadOnly(True)
        browse_button = QPushButton("瀏覽..."); browse_button.clicked.connect(self.select_folder)
        h_layout.addWidget(QLabel("選擇檔案資料夾：")); h_layout.addWidget(self.folder_path_edit); h_layout.addWidget(browse_button)
        layout.addLayout(h_layout); self.start_button = QPushButton("開始處理"); self.start_button.setEnabled(False)
        self.start_button.clicked.connect(self.start_processing); layout.addWidget(self.start_button)
        layout.addWidget(QLabel("處理日誌：")); self.log_display = QTextEdit(); self.log_display.setReadOnly(True)
        layout.addWidget(self.log_display)
    def select_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, "選擇資料夾")
        if folder_path: self.folder_path_edit.setText(folder_path); self.start_button.setEnabled(True); self.log_display.append(f"已選擇資料夾：{folder_path}")
    def start_processing(self): self.start_task(run_processor_task, self.folder_path_edit.text(), os.path.basename(self.folder_path_edit.text()))

class ConverterUI(BaseUI):
    def __init__(self): super().__init__(); self.init_ui()
    def init_ui(self):
        layout = QVBoxLayout(self)
        self.file_path_edit = self._create_file_selector(layout, "選擇 SRT 檔案：", "SRT Files (*.srt)")
        self.start_button = QPushButton("開始轉換"); self.start_button.setEnabled(False)
        self.start_button.clicked.connect(self.start_conversion); layout.addWidget(self.start_button)
        layout.addWidget(QLabel("處理日誌：")); self.log_display = QTextEdit(); self.log_display.setReadOnly(True)
        layout.addWidget(self.log_display)
    def _create_file_selector(self, layout, label, file_filter):
        layout.addWidget(QLabel(label)); h_layout = QHBoxLayout(); line_edit = QLineEdit(); line_edit.setReadOnly(True)
        browse_button = QPushButton("瀏覽..."); browse_button.clicked.connect(lambda: self._select_file(line_edit, file_filter))
        h_layout.addWidget(line_edit); h_layout.addWidget(browse_button); layout.addLayout(h_layout); return line_edit
    def _select_file(self, line_edit, file_filter):
        file_name, _ = QFileDialog.getOpenFileName(self, "選擇檔案", "", file_filter)
        if file_name:
            line_edit.setText(file_name); self.log_display.append(f"已選擇: {os.path.basename(file_name)}")
            self.start_button.setEnabled(True)
    def start_conversion(self): self.start_task(run_converter_task, self.file_path_edit.text())
        
class DocxFormatterUI(BaseUI):
    def __init__(self): super().__init__(); self.init_ui()
    def init_ui(self):
        layout = QVBoxLayout(self)
        self.file_path_edit = self._create_file_selector(layout, "選擇 Excel 檔案：", "Excel Files (*.xlsx)")
        self.start_button = QPushButton("產生校對用 Word 文件"); self.start_button.setEnabled(False)
        self.start_button.clicked.connect(self.start_formatting); layout.addWidget(self.start_button)
        layout.addWidget(QLabel("處理日誌：")); self.log_display = QTextEdit(); self.log_display.setReadOnly(True)
        layout.addWidget(self.log_display)
    def _create_file_selector(self, layout, label, file_filter):
        return ConverterUI._create_file_selector(self, layout, label, file_filter)
    def _select_file(self, line_edit, file_filter):
        ConverterUI._select_file(self, line_edit, file_filter)
    def start_formatting(self): self.start_task(run_docx_formatter_task, self.file_path_edit.text())

class MainApp(QMainWindow):
    def __init__(self):
        super().__init__(); self.setWindowTitle("多功能檔案處理工具 v2.8"); self.tabs = QTabWidget(); self.setCentralWidget(self.tabs)
        self.tabs.addTab(AlignerUI(), "1. 文本對齊"); self.tabs.addTab(ConverterUI(), "2. 加入時間軸到 Excel")
        self.tabs.addTab(ProcessorUI(), "3. 逐字稿分段處理"); self.tabs.addTab(DocxFormatterUI(), "4. 產生校對用Word檔")
        self.setGeometry(100, 100, 750, 600)

# ===================================================================
#
#  第四部分：程式執行入口
#
# ===================================================================
if __name__ == '__main__':
    app = QApplication(sys.argv)
    main_window = MainApp()
    main_window.show()
    sys.exit(app.exec())