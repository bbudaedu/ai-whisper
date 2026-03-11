import os
import sys
import re
import math
import difflib
import pandas as pd
import openpyxl
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl.utils import get_column_letter

import logging
import json
import time

from pipeline.api_client import ResilientAPIClient

# 設定 logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 載入設定檔 (config.json)
CONFIG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")
config_data = {}
if os.path.exists(CONFIG_FILE):
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            config_data = json.load(f)
    except Exception as e:
        logger.error(f"無法讀取 config.json: {e}")

# 中轉 API 設定
API_BASE_URL = config_data.get("api_base_url", "http://192.168.100.201:8045/v1/chat/completions")
API_KEY = config_data.get("api_key", "sk-8f3999c2452d4124835ffaff469e22af")
PUNCTUATION_MODEL = config_data.get("punctuation_model", "gemini-2.5-flash")
PUNCTUATION_CHUNK_SIZE = config_data.get("punct_chunk_size", 120)

def simple_autofit_columns(output_path):
    try:
        wb = openpyxl.load_workbook(output_path)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for col in ws.columns:
                max_length = 0
                column_letter = get_column_letter(col[0].column)
                header_value = ws[1][col[0].column - 1].value
                if header_value:
                    max_length = len(str(header_value)) * 1.2

                for cell in col:
                    try:
                        if cell.value:
                            cell_len = sum(2 if '\u4e00' <= char <= '\u9fff' else 1 for char in str(cell.value))
                            if cell_len > max_length:
                                max_length = cell_len
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column_letter].width = adjusted_width
        wb.save(output_path)
        logger.info(f"✅ Excel 欄寬自動調整完成: {output_path}")
    except Exception as e:
        logger.error(f"⚠️ 自動調整欄寬時發生錯誤: {e}")

def get_similarity_ratio(s1, s2):
    return difflib.SequenceMatcher(None, s1, s2, autojunk=False).ratio()

def align_sequences(seq1, seq2, gap_penalty=-0.5):
    n, m = len(seq1), len(seq2)
    dp = [[(0, '')] * (m + 1) for _ in range(n + 1)]
    for i in range(1, n + 1): dp[i][0] = (i * gap_penalty, 'up')
    for j in range(1, m + 1): dp[0][j] = (j * gap_penalty, 'left')
    
    for i in range(1, n + 1):
        for j in range(1, m + 1):
            match = dp[i-1][j-1][0] + (2 * get_similarity_ratio(seq1[i-1], seq2[j-1]) - 1)
            scores = {'diag': match, 'up': dp[i-1][j][0] + gap_penalty, 'left': dp[i][j-1][0] + gap_penalty}
            best_move = max(scores, key=scores.get)
            dp[i][j] = (scores[best_move], best_move)

    align1, align2 = [], []
    i, j = n, m
    while i > 0 or j > 0:
        move = dp[i][j][1]
        if move == 'diag': align1.append(seq1[i-1]); align2.append(seq2[j-1]); i-=1; j-=1
        elif move == 'up': align1.append(seq1[i-1]); align2.append(""); i-=1
        else: align1.append(""); align2.append(seq2[j-1]); j-=1
    return align1[::-1], align2[::-1]

def read_srt_lines(srt_path):
    lines = []
    current_text = []
    with open(srt_path, 'r', encoding='utf-8-sig') as f:
        for line in f:
            line = line.strip()
            if not line:
                if current_text:
                    lines.append(" ".join(current_text))
                    current_text = []
            elif not re.match(r'^\d+$', line) and '-->' not in line:
                current_text.append(line)
        if current_text:
            lines.append(" ".join(current_text))
    return lines

def srt_time_to_seconds(time_str):
    if not time_str or time_str.count(':') != 2 or ',' not in time_str: return 0.0
    parts = time_str.split(',')
    h_m_s = parts[0].split(':')
    return int(h_m_s[0])*3600 + int(h_m_s[1])*60 + int(h_m_s[2]) + int(parts[1])/1000.0

# 建立統一 API 客戶端（含指數退避 + Circuit Breaker）
_api_client = ResilientAPIClient(
    api_base_url=API_BASE_URL,
    api_key=API_KEY,
    model=PUNCTUATION_MODEL,
)


def call_gemini_api(prompt, max_retries=3):
    """呼叫中轉 API 進行標點與分段（透過 ResilientAPIClient）"""
    return _api_client.call(prompt)

def format_text_with_ai(sentences):
    """將句子分批送給 AI 加上標點與分段"""
    logger.info(f"開始 AI 自動排版與標點，共 {len(sentences)} 句")
    chunks = [sentences[i:i+PUNCTUATION_CHUNK_SIZE] for i in range(0, len(sentences), PUNCTUATION_CHUNK_SIZE)]
    total_chunks = len(chunks)
    
    formatted_paragraphs = []
    
    for i, chunk in enumerate(chunks, 1):
        logger.info(f"正在處理 AI 標點排版批次 {i}/{total_chunks} ({len(chunk)} 句)")
        chunk_text = " ".join(chunk).strip()
        
        prompt = f"""你是一位專業的佛學文獻編輯與排版專家。
我會給你一段由語音辨識產生的佛學講座逐字稿（純文字，無標點符號）。

你的任務只有兩個：
1. 【精準標點】：請根據語氣和佛學專有名詞，為這段文字加上正確、通順的中文全形標點符號（，、。！？；：「」『』等）。
2. 【邏輯分段】：請根據講者的語境、主題切換以及呼吸停頓，適當地將長篇大論切割成易於閱讀的段落（Paragraphs）。

【嚴格遵守以下規則 - 違反將導致任務失敗】：
- 絕對禁止增減、替換、改寫講者的任何一個字！即使你認為講者有口誤，也必須保留原字。
- 請不要加上任何你的解釋、反饋或開場白（例如：不要說「好的」或「修改後的文本如下」），直接輸出處理完的最終內文。
- 請在遇到明顯的儀軌或開示段落時換行，例如：「請諸位大德一起合掌，恭念三稱本師聖號及開經偈」這類引言後應當換行。

請開始處理以下逐字稿段落：
{chunk_text}"""
        
        result = call_gemini_api(prompt)
        if result:
            formatted_paragraphs.append(result.strip())
        else:
            logger.warning(f"批次 {i} 標點處理失敗，退回原始文字。")
            formatted_paragraphs.append(chunk_text)
            
        if i < total_chunks:
            time.sleep(2)
            
    return "\n\n".join(formatted_paragraphs)

def generate_excel_and_docx(episode_dir, base_name):
    """
    執行對齊 -> Excel 生成 -> Docx 生成的一條龍處理。
    episode_dir: 存放檔案的資料夾路徑 (例如: /mnt/nas/Whisper_auto_rum/T097V/T097V011)
    base_name: 檔案基底名稱 (例如: 佛教公案選集 簡豐文居士 011__keoeMspJAcU)
    """
    logger.info(f"{'='*60}")
    logger.info(f"開始排版與文件生成作業: {base_name}")
    logger.info(f"{'='*60}")

    whisper_srt = os.path.join(episode_dir, f"{base_name}.srt")
    gemini_srt = os.path.join(episode_dir, f"{base_name}_proofread.srt")
    excel_path = os.path.join(episode_dir, f"{base_name}.xlsx")
    
    if not os.path.exists(whisper_srt):
        logger.error(f"找不到 Whisper SRT 檔案: {whisper_srt}")
        return False
        
    has_gemini = os.path.exists(gemini_srt)
    if not has_gemini:
        logger.warning(f"找不到 Gemini 修正版 SRT ({gemini_srt})，將僅使用 Whisper 原始字幕。")
        gemini_srt = whisper_srt

    # 1. 解析 SRT 時間軸
    logger.info("解析 SRT 時間軸...")
    with open(whisper_srt, 'r', encoding='utf-8-sig') as f:
        content = f.read()
    time_data = re.findall(r'(\d+)\n(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})\n(.*?)(?=\n\n|\Z)', content, re.DOTALL)
    
    df_timeline = pd.DataFrame(time_data, columns=['序號', '開始時間', '結束時間', '文字'])
    df_timeline['文字'] = df_timeline['文字'].str.replace('\n', ' ', regex=False)

    # 2. 文本對齊
    logger.info("執行 Whisper 與 Gemini 文本對齊...")
    lines_whisper = read_srt_lines(whisper_srt)
    lines_gemini = read_srt_lines(gemini_srt)
    
    aligned_whisper, aligned_gemini = align_sequences(lines_whisper, lines_gemini)
    df_aligned = pd.DataFrame({'whisper': aligned_whisper, 'gemini': aligned_gemini})
    df_aligned = df_aligned[df_aligned['whisper'] != ''].copy()
    df_aligned.loc[df_aligned['gemini'] == '', 'gemini'] = df_aligned['whisper']

    final_gemini_texts = df_aligned['gemini'].astype(str).str.strip().tolist()
    
    # 建立 Excel Writer
    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        # Sheet: 時間軸
        df_timeline.to_excel(writer, sheet_name='時間軸', index=False)
        # Sheet: 文本校對
        df_aligned.to_excel(writer, sheet_name='文本校對', index=False)
        
        # 3. 生成「給學員校對」分段
        logger.info("產生「給學員校對」分組時間區塊...")
        raw_start_times = df_timeline['開始時間'].tolist()
        raw_end_times = df_timeline['結束時間'].tolist()
        
        min_len = min(len(final_gemini_texts), len(raw_start_times))
        items_with_time = []
        for i in range(min_len):
            items_with_time.append({
                'num': i + 1,
                'text': f"{i + 1} {final_gemini_texts[i]}",
                'start_s': srt_time_to_seconds(raw_start_times[i]),
                'end_s': srt_time_to_seconds(raw_end_times[i]),
                'start_str': raw_start_times[i],
                'end_str': raw_end_times[i]
            })

        total_audio_duration_seconds = items_with_time[-1]['end_s'] if items_with_time else 0
        segment_duration_seconds = 30 * 60
        total_parts = math.ceil(total_audio_duration_seconds / segment_duration_seconds) if total_audio_duration_seconds > 0 else 1
        
        rows_for_excel_output = []
        part_number = 1
        current_part_items = []
        
        for item in items_with_time:
            segment_boundary_s = part_number * segment_duration_seconds
            if item['start_s'] >= segment_boundary_s and current_part_items:
                start_sentence_num, end_sentence_num = current_part_items[0]['num'], current_part_items[-1]['num']
                header_line1 = f"{base_name} Part {part_number} of {total_parts}"
                header_line2 = f"{start_sentence_num}-{end_sentence_num}句"
                header_line3 = f"{current_part_items[0]['start_str']} --> {current_part_items[0]['end_str']}"
                rows_for_excel_output.extend([[header_line1], [header_line2], [header_line3]])
                for part_item in current_part_items: rows_for_excel_output.append([part_item['text']])
                rows_for_excel_output.append([""])
                part_number += 1
                current_part_items = []
            current_part_items.append(item)
            
        if current_part_items:
            start_sentence_num, end_sentence_num = current_part_items[0]['num'], current_part_items[-1]['num']
            header_line1 = f"{base_name} Part {part_number} of {total_parts}"
            header_line2 = f"{start_sentence_num}-{end_sentence_num}句"
            header_line3 = f"{current_part_items[0]['start_str']} --> {current_part_items[0]['end_str']}"
            rows_for_excel_output.extend([[header_line1], [header_line2], [header_line3]])
            for part_item in current_part_items: rows_for_excel_output.append([part_item['text']])

        df_student = pd.DataFrame(rows_for_excel_output, columns=['給學員校對用'])
        df_student.to_excel(writer, sheet_name='給學員校對', index=False)
        logger.info("✅ Excel 核心資料寫入完成")

    # 調整 Excel 欄寬自動適應
    simple_autofit_columns(excel_path)

    # 4. 生成 Word (Docx) 給學員校對
    logger.info("生成 Docx 文件...")
    text_content = "\n".join([str(rows[0]) for rows in rows_for_excel_output if rows[0] is not None])
    
    def create_docx(out_path):
        document = docx.Document()
        style = document.styles['Normal']
        font = style.font
        p_format = style.paragraph_format
        
        font.name = '標楷體'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        font.size = Pt(16)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_format.line_spacing = Pt(24)

        section = document.sections[0]
        section.top_margin = Cm(1.27); section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27); section.right_margin = Cm(1.27)
        section.header_distance = Cm(0.6); section.footer_distance = Cm(0.6)
        
        footer = section.footer
        p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_footer.add_run()
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
        
        sectPr = section._sectPr
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), str(Cm(1).twips))
        cols.set(qn('w:sep'), 'true')
        sectPr.append(cols)
        
        splitter = f'{base_name} Part '
        parts_raw = text_content.split(splitter)[1:]
        
        is_first_part = True
        for part_content in parts_raw:
            if not is_first_part:
                document.add_section(WD_SECTION_START.NEW_PAGE)
                new_section = document.sections[-1]
                new_section.top_margin = Cm(1.27); new_section.bottom_margin = Cm(1.27)
                new_section.left_margin = Cm(1.27); new_section.right_margin = Cm(1.27)
                new_section.header_distance = Cm(0.6); new_section.footer_distance = Cm(0.6)
                sectPr = new_section._sectPr
                cols = OxmlElement('w:cols')
                cols.set(qn('w:num'), '2'); cols.set(qn('w:space'), str(Cm(1).twips)); cols.set(qn('w:sep'), 'true')
                sectPr.append(cols)

            lines = part_content.strip().split('\n')
            header_line = splitter + lines[0]
            
            doc_para = document.add_paragraph()
            doc_para.text = header_line

            for line in lines[1:]:
                doc_para = document.add_paragraph()
                doc_para.text = line
            is_first_part = False
            
        document.save(out_path)

    title_part = base_name.split('__')[0]
    match = re.search(r'佛教公案選集\s*簡豐文居士\s*(\d+)', title_part)
    if match:
        ep_num = match.group(1)
        prefix = f"佛教公案選集{ep_num}_簡豐文居士"
    else:
        prefix = title_part.replace(" ", "_")[:25]

    docx_student = os.path.join(episode_dir, f"{prefix}給學員校對.docx")
    docx_ai = os.path.join(episode_dir, f"{prefix}校對文本.docx")

    create_docx(docx_student)
    logger.info(f"✅ Docx (給學員校對) 建立成功: {docx_student}")

    # 5. 生成 Word (Docx) 供最終閱讀 (AI 排版與標點)
    logger.info("正在生成 AI 標點排版校對文本...")
    try:
        ai_formatted_text = format_text_with_ai(final_gemini_texts)
    except Exception as e:
        logger.error(f"AI 排版失敗，將使用沒有標點的原始文字: {e}")
        ai_formatted_text = "\n".join(final_gemini_texts)

    def create_ai_docx(out_path, full_text):
        document = docx.Document()
        style = document.styles['Normal']
        font = style.font
        p_format = style.paragraph_format
        
        font.name = '標楷體'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        font.size = Pt(16)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(12)
        p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_format.line_spacing = Pt(24)

        section = document.sections[0]
        section.top_margin = Cm(1.27); section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27); section.right_margin = Cm(1.27)

        title_para = document.add_paragraph(base_name)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.bold = True
        title_para.runs[0].font.size = Pt(18)

        paragraphs = full_text.split('\n')
        for p in paragraphs:
            if p.strip():
                doc_para = document.add_paragraph(p.strip())
                doc_para.paragraph_format.first_line_indent = Pt(32)

        document.save(out_path)

    create_ai_docx(docx_ai, ai_formatted_text)
    logger.info(f"✅ Docx (AI 校對文本) 建立成功: {docx_ai}")
    
    return [excel_path, docx_student, docx_ai]

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python auto_postprocess.py <episode_dir> <base_name>")
        sys.exit(1)
    
    ep_dir = sys.argv[1]
    b_name = sys.argv[2]
    generate_excel_and_docx(ep_dir, b_name)
